"""
Скрипт для пересоздания шаблона template.docx с правильной структурой docxtpl.

Запуск: poetry run python src/teacherfactory/rebuild_template.py
"""

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path


PROJECT_ROOT = Path(__file__).parent.parent.parent
OUTPUT_PATH = PROJECT_ROOT / "template_fixed.docx"


def build_template():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # === Основная таблица ===
    table = doc.add_table(rows=29, cols=3)
    table.style = 'Table Grid'

    def merge_and_set(row_idx, text, merge_cols=True):
        row = table.rows[row_idx]
        if merge_cols:
            row.cells[0].merge(row.cells[2])
        row.cells[0].text = text

    def set_row(row_idx, label, value_template):
        table.rows[row_idx].cells[0].text = label
        table.rows[row_idx].cells[1].merge(table.rows[row_idx].cells[2])
        table.rows[row_idx].cells[1].text = value_template

    # Строка 0 — заголовок (объединённая)
    merge_and_set(0, "Занятие №{{ lesson_number }} \nДата: {{ date }}\nГруппа: {{ group_name }}\nКурс: {{ course_number }}\nКол-во студентов: {{ students_count }}")

    # Строки 1-7 — основная информация
    set_row(1, "Дисциплина / МДК", "{{ discipline }}")
    set_row(2, "Специальность", "{{ specialty }}")
    set_row(3, "Тема занятия", "{{ lesson_topic }}")
    set_row(4, "Тип занятия", "{{ lesson_type }}")
    set_row(5, "Вид занятия", "{{ lesson_kind }}")
    set_row(6, "Продолжительность", "{{ duration }} мин")
    set_row(7, "Преподаватель", "{{ teacher_name }}")

    # Строка 8 — цель
    set_row(8, "Цель занятия", "{{ goal }}")

    # Строки 9-10 — образовательная задача
    set_row(9, "Задачи занятия", "Образовательная задача")
    set_row(10, "", "{{ task_educational }}")

    # Строки 11-12 — развивающая
    set_row(11, "Задачи занятия", "Развивающая задача")
    set_row(12, "", "{{ task_developmental }}")

    # Строки 13-14 — воспитательная
    set_row(13, "Задачи занятия", "Воспитательная задача")
    set_row(14, "", "{{ task_upbringing }}")

    # Строки 15-16 — педагогическая
    set_row(15, "Задачи занятия", "Педагогическая задача (методическая цель)")
    set_row(16, "", "{{ task_pedagogical }}")

    # Строки 17-18 — компетенции
    set_row(17, "Общие компетенции (ОК)", "{{ competencies_ok }}")
    set_row(18, "Профессиональные компетенции (ПК)", "{{ competencies_pk }}")

    # Строки 19-25 — планируемые результаты
    row19 = table.rows[19]
    row19.cells[0].text = "Планируемые результаты"
    row19.cells[1].text = "Индикатор"
    row19.cells[2].text = "Форма контроля"

    set_label_and_values = [
        (20, "Знания", "{{ knowledge_indicator }}", "{{ knowledge_control }}"),
        (21, "", "", ""),
        (22, "Умения", "{{ skill_indicator }}", "{{ skill_control }}"),
        (23, "", "", ""),
        (24, "Навыки", "{{ ability_indicator }}", "{{ ability_control }}"),
        (25, "", "", ""),
    ]
    for row_idx, label, indicator, control in set_label_and_values:
        if label:
            table.rows[row_idx].cells[0].text = label
            table.rows[row_idx].cells[1].text = indicator
            table.rows[row_idx].cells[2].text = control

    # Строки 26-28 — ресурсы
    set_row(26, "Литература", "{{ literature }}")
    set_row(27, "Интернет-ресурсы", "{{ internet_resources }}")
    set_row(28, "Другие материалы", "{{ other_materials }}")

    # === Таблица хода занятия ===
    doc.add_paragraph()  # пустая строка
    doc.add_paragraph("Ход занятия")

    lesson_table = doc.add_table(rows=3, cols=7)
    lesson_table.style = 'Table Grid'

    # Заголовок
    headers = ["№ п/п", "Этап", "Время, мин.", "Деятельность преподавателя",
               "Деятельность студентов", "Методы и приёмы обучения", "Результат"]
    for i, h in enumerate(headers):
        lesson_table.rows[0].cells[i].text = h

    # for и endfor в РАЗНЫХ строках — docxtpl удаляет обе строки-маркеры,
    # строку 1 дублирует для каждого шага, строку 2 убирает целиком
    template_row = lesson_table.rows[1]
    template_row.cells[0].text = "{%tr for step in lesson_structure %}{{ step.number }}"
    template_row.cells[1].text = "{{ step.stage }}"
    template_row.cells[2].text = "{{ step.time }}"
    template_row.cells[3].text = "{{ step.teacher }}"
    template_row.cells[4].text = "{{ step.student }}"
    template_row.cells[5].text = "{{ step.methods }}"
    template_row.cells[6].text = "{{ step.result }}"

    lesson_table.rows[2].cells[0].text = "{%tr endfor %}"

    doc.save(str(OUTPUT_PATH))
    print(f"Шаблон пересоздан: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_template()
