"""
Индексатор документов для RAG.

Загружает PDF и DOCX из папки docs/, нарезает на чанки с учётом структуры
российских нормативных документов (разделы, таблицы, абзацы),
создаёт FAISS-индекс (семантический) и BM25-индекс (ключевые слова).

Запуск: poetry run python -m teacherfactory.indexer
"""

import logging
from pathlib import Path

from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

from teacherfactory.config import CONFIG
from teacherfactory.embeddings import get_embeddings
from teacherfactory.paths import DOCS_DIR, INDEX_DIR, SKIPPED_DOC_PATH
from teacherfactory.retrieval import save_bm25

log = logging.getLogger(__name__)

# Разделители в порядке приоритета: сначала крупные структурные границы,
# потом мелкие. Строки (\n) идут раньше точек — это помогает не резать
# таблицы компетенций посередине строки.
_SEPARATORS = [
    "\n\n\n",  # крупные разделы
    "\n\n",  # абзацы
    "\n",  # строки (важно для таблиц)
    ". ",  # предложения
    ", ",  # перечисления
    " ",  # слова
    "",  # символы (крайний случай)
]


def _load_pdf(pdf_path: Path) -> list[Document]:
    """Загрузить PDF через pdfplumber — сохраняет структуру таблиц компетенций."""
    import pdfplumber

    docs: list[Document] = []
    try:
        with pdfplumber.open(str(pdf_path)) as pdf:
            for i, page in enumerate(pdf.pages):
                parts: list[str] = []

                # Сначала извлекаем таблицы — они содержат коды ОК/ПК
                tables = page.extract_tables()
                for table in tables:
                    for row in table:
                        if not row:
                            continue
                        cells = [str(c).strip() for c in row if c and str(c).strip()]
                        if cells:
                            parts.append("\t".join(cells))

                # Затем основной текст страницы
                text = page.extract_text(x_tolerance=3, y_tolerance=3) or ""
                if text.strip():
                    parts.append(text.strip())

                full_text = "\n".join(parts).strip()
                if full_text:
                    docs.append(
                        Document(
                            page_content=full_text,
                            metadata={"source": str(pdf_path), "page": i},
                        )
                    )
    except Exception as e:
        log.warning("pdfplumber не смог прочитать %s: %s — пробую PyPDF", pdf_path.name, e)
        docs = _load_pdf_fallback(pdf_path)

    return docs


def _load_pdf_fallback(pdf_path: Path) -> list[Document]:
    """Запасной вариант — PyPDF (хуже обрабатывает таблицы)."""
    from langchain_community.document_loaders import PyPDFLoader

    try:
        loader = PyPDFLoader(str(pdf_path))
        return loader.load()
    except Exception as e:
        log.error("PyPDF fallback failed for %s: %s", pdf_path.name, e)
        return []


def _load_docx(docx_path: Path) -> list[Document]:
    """Загрузить DOCX с сохранением структуры таблиц компетенций."""
    try:
        from docx import Document as DocxDoc
    except ImportError:
        log.warning("python-docx не установлен, пропускаю %s", docx_path.name)
        return []

    try:
        word_doc = DocxDoc(str(docx_path))
        parts: list[str] = []

        for para in word_doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        for table in word_doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    # Убираем дубли (merged cells дублируют содержимое)
                    unique_cells = list(dict.fromkeys(cells))
                    parts.append("\t".join(unique_cells))

        full_text = "\n".join(parts).strip()
        if full_text:
            return [
                Document(
                    page_content=full_text,
                    metadata={"source": str(docx_path), "page": 0},
                )
            ]
    except Exception as e:
        log.warning("Не удалось прочитать %s: %s", docx_path.name, e)

    return []


def build_index() -> dict:
    """
    Построить FAISS + BM25 индексы из PDF и DOCX в папке docs/.

    Возвращает словарь со статистикой:
        {
            "chunks": int,         # сколько чанков попало в индекс
            "documents": int,      # сколько исходных PDF/DOCX
            "skipped_doc": list,   # .doc файлы, пропущенные как неподдерживаемые
            "embeddings": str,     # имя класса эмбеддингов
            "dim": int,            # размерность вектора
        }
    """
    all_docs: list[Document] = []
    skipped_doc: list[str] = []

    pdf_files = sorted(DOCS_DIR.glob("*.pdf"))
    docx_files = sorted(DOCS_DIR.glob("*.docx"))
    doc_files = sorted(DOCS_DIR.glob("*.doc"))

    if not pdf_files and not docx_files:
        raise FileNotFoundError(f"В папке {DOCS_DIR} не найдено ни одного .pdf или .docx файла.")

    for pdf_path in pdf_files:
        log.info("  PDF: %s", pdf_path.name)
        pages = _load_pdf(pdf_path)
        all_docs.extend(pages)
        log.info("    → %d страниц", len(pages))

    for docx_path in docx_files:
        log.info("  DOCX: %s", docx_path.name)
        pages = _load_docx(docx_path)
        all_docs.extend(pages)
        log.info("    → %d фрагментов", len(pages))

    for doc_path in doc_files:
        log.warning("  ⚠ .doc пропущен (конвертируй в .docx или .pdf): %s", doc_path.name)
        skipped_doc.append(doc_path.name)

    if skipped_doc:
        log.warning(
            "Пропущено %d .doc файлов: %s. Сконвертируй их в .docx через Word или LibreOffice.",
            len(skipped_doc),
            skipped_doc,
        )

    log.info(
        "Загружено %d фрагментов из %d PDF + %d DOCX файлов",
        len(all_docs),
        len(pdf_files),
        len(docx_files),
    )

    if not all_docs:
        raise RuntimeError(
            f"Не удалось извлечь текст ни из одного документа в {DOCS_DIR}. "
            "Возможно, файлы повреждены или защищены."
        )

    chunk_size = CONFIG["rag"]["chunk_size"]
    chunk_overlap = CONFIG["rag"]["chunk_overlap"]
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=_SEPARATORS,
    )
    chunks = splitter.split_documents(all_docs)
    log.info("Нарезано %d чанков (size=%d, overlap=%d)", len(chunks), chunk_size, chunk_overlap)

    INDEX_DIR.mkdir(parents=True, exist_ok=True)

    embeddings = get_embeddings()
    log.info("Создаю эмбеддинги: %s", type(embeddings).__name__)
    db = FAISS.from_documents(chunks, embeddings)
    db.save_local(str(INDEX_DIR))
    log.info("FAISS-индекс сохранён в %s", INDEX_DIR)

    save_bm25(chunks)
    log.info("BM25-индекс сохранён (JSON, %d чанков)", len(chunks))

    if skipped_doc:
        SKIPPED_DOC_PATH.write_text("\n".join(skipped_doc), encoding="utf-8")
    elif SKIPPED_DOC_PATH.exists():
        SKIPPED_DOC_PATH.unlink()

    log.info("--- Тестовый поиск: 'профессиональные компетенции ПК' ---")
    results = db.similarity_search("профессиональные компетенции ПК", k=3)
    for i, doc in enumerate(results, 1):
        source = Path(doc.metadata.get("source", "?")).name
        log.info(
            "[%d] %s (стр. %s): %s...",
            i,
            source,
            doc.metadata.get("page", "?"),
            doc.page_content[:150],
        )

    return {
        "chunks": len(chunks),
        "documents": len(pdf_files) + len(docx_files),
        "skipped_doc": skipped_doc,
        "embeddings": type(embeddings).__name__,
        "dim": db.index.d,
    }


def main() -> None:
    logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s %(message)s")
    build_index()


if __name__ == "__main__":
    main()
