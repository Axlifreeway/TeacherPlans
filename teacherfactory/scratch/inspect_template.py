import sys
import docx

doc = docx.Document("docs/template_aleshkin_v2.docx")
sys.stderr.write(f"tables: {len(doc.tables)}\n")
for ti, t in enumerate(doc.tables):
    sys.stderr.write(f"\n=== TABLE {ti}: {len(t.rows)} rows, {len(t.rows[0].cells) if t.rows else 0} cols ===\n")
    for ri, row in enumerate(t.rows[:4]):
        for ci, cell in enumerate(row.cells):
            txt = cell.text.replace("\n", " | ")
            sys.stderr.write(f"  R{ri}C{ci}: {txt[:80]!r}\n")
