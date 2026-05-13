import docx
import sys
import json
from pathlib import Path

doc_path = Path('docs/16.04.26 Алёшкин АА Технологическая карта SQL.docx')
doc = docx.Document(doc_path)
out = {
    'paragraphs': [p.text for p in doc.paragraphs if p.text.strip()],
    'tables': [[[c.text.strip() for c in row.cells] for row in t.rows] for t in doc.tables]
}
with open('scratch_docx.json', 'w', encoding='utf-8') as f:
    json.dump(out, f, ensure_ascii=False, indent=2)
