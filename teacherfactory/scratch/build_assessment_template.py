"""
Минимальный шаблон оценочных материалов (КОС/ФОС) → templates/template_assessment.docx.

Это «фундамент»: структура полей соответствует AssessmentMaterials,
визуальная вёрстка минимальная. Когда появится финальная вёрстка от методиста,
заменим этот скрипт на копию реального docx с jinja-тегами (как для карты).

Запуск: poetry run python scratch/build_assessment_template.py
"""

from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt

DST = Path("templates/template_assessment.docx")


def build() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    # Шапка
    p = doc.add_paragraph()
    p.add_run("ОЦЕНОЧНЫЕ МАТЕРИАЛЫ ПО ДИСЦИПЛИНЕ").bold = True

    doc.add_paragraph("Дисциплина: {{ discipline }}")
    doc.add_paragraph("Специальность: {{ specialty }}")
    doc.add_paragraph("Семестр: {{ semester }}")
    doc.add_paragraph("Форма промежуточной аттестации: {{ final_form }}")

    doc.add_heading("Система оценивания", level=2)
    doc.add_paragraph("{{ grading_system }}")

    # Таблица «Паспорт ОМ»: 3 строки (header, for-marker, body, endfor-marker).
    doc.add_heading("Паспорт оценочных материалов", level=2)
    t = doc.add_table(rows=4, cols=4)
    t.style = "Table Grid"

    headers = ["Темы дисциплины", "Оценочные материалы", "Код компетенции", "Критерии оценивания"]
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    # 3-рядный паттерн loop по docxtpl
    t.rows[1].cells[0].text = "{%tr for row in passport_rows %}"
    t.rows[2].cells[0].text = "{{ row.topic }}"
    t.rows[2].cells[1].text = "{{ row.tool }}"
    t.rows[2].cells[2].text = "{{ row.competencies }}"
    t.rows[2].cells[3].text = "{{ row.criteria }}"
    t.rows[3].cells[0].text = "{%tr endfor %}"

    # Таблица «Компетенции — количество заданий»
    doc.add_heading("Распределение по компетенциям", level=2)
    t2 = doc.add_table(rows=4, cols=3)
    t2.style = "Table Grid"
    for i, h in enumerate(["Код", "Наименование", "Количество заданий"]):
        cell = t2.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    t2.rows[1].cells[0].text = "{%tr for s in competency_summary %}"
    t2.rows[2].cells[0].text = "{{ s.code }}"
    t2.rows[2].cells[1].text = "{{ s.name }}"
    t2.rows[2].cells[2].text = "{{ s.items_count }}"
    t2.rows[3].cells[0].text = "{%tr endfor %}"

    doc.add_paragraph()
    doc.add_paragraph("ВСЕГО заданий: {{ competency_total }}")

    # Тестовые задания
    doc.add_heading("Типовые тестовые задания", level=2)
    t3 = doc.add_table(rows=4, cols=4)
    t3.style = "Table Grid"
    for i, h in enumerate(["Код", "№", "Инструкция и задание", "Ключ"]):
        cell = t3.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    t3.rows[1].cells[0].text = "{%tr for item in test_items %}"
    t3.rows[2].cells[0].text = "{{ item.competency_code }}"
    t3.rows[2].cells[1].text = "{{ item.number }}"
    t3.rows[2].cells[2].text = "{{ item.instruction_and_task }}"
    t3.rows[2].cells[3].text = "{{ item.answer_key }}"
    t3.rows[3].cells[0].text = "{%tr endfor %}"

    DST.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(DST))
    print(f"Шаблон оценочных материалов: {DST}")


if __name__ == "__main__":
    build()
