import sys
import docx

doc = docx.Document("docs/16.04.26 Алёшкин АА Технологическая карта SQL.docx")
for i, p in enumerate(doc.paragraphs):
    if "endfor" in p.text or "{%" in p.text:
        sys.stderr.write(f"PARA {i}: {p.text!r}\n")
for ti, t in enumerate(doc.tables):
    for ri, row in enumerate(t.rows):
        for ci, cell in enumerate(row.cells):
            if "endfor" in cell.text or "{%" in cell.text:
                sys.stderr.write(f"TABLE {ti} R{ri} C{ci}: {cell.text!r}\n")
sys.stderr.write("done\n")
