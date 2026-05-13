import sys
import docx

doc = docx.Document("scratch/test_output.docx")
sys.stderr.write(f"paragraphs: {len(doc.paragraphs)}, tables: {len(doc.tables)}\n")

# Все ли значения подставились (не осталось ли {{ ... }})?
import re

for i, p in enumerate(doc.paragraphs):
    if "{{" in p.text or "{%" in p.text:
        sys.stderr.write(f"PARA {i} HAS JINJA: {p.text!r}\n")
for ti, t in enumerate(doc.tables):
    for ri, row in enumerate(t.rows):
        for ci, cell in enumerate(row.cells):
            if "{{" in cell.text or "{%" in cell.text:
                sys.stderr.write(f"T{ti}R{ri}C{ci} HAS JINJA: {cell.text!r}\n")
sys.stderr.write("\n=== Sample content (paragraphs 0-30) ===\n")
for i, p in enumerate(doc.paragraphs[:30]):
    if p.text.strip():
        sys.stderr.write(f"  {i}: {p.text[:100]}\n")

sys.stderr.write("\n=== Table sizes ===\n")
for ti, t in enumerate(doc.tables):
    sys.stderr.write(f"  table {ti}: {len(t.rows)} rows\n")

# Проверим, есть ли ожидаемые подстановки
text = " ".join(p.text for p in doc.paragraphs)
text += " " + " ".join(
    cell.text for t in doc.tables for row in t.rows for cell in row.cells
)
expected_substrings = [
    "SELECT с WHERE",  # lesson_topic
    "Алёшкин А.А.",  # teacher_name
    "ОИС-25-1",  # group_name
    "ОК 01",  # competency code
    "Кнут",  # epigraph_author
    "Организационный",  # lesson step name
]
sys.stderr.write("\n=== Substitution check ===\n")
for s in expected_substrings:
    sys.stderr.write(f"  {s!r}: {'OK' if s in text else 'MISSING'}\n")
