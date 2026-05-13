import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

import docx

from teacherfactory.documents.lesson_card import LESSON_CARD
from teacherfactory.render import render_document
from tests.conftest import build_stub_card

card = build_stub_card()
out = Path("scratch/full_stub_test.docx")
render_document(LESSON_CARD, card, out)
sys.stderr.write(f"OK render: {out} -> {out.stat().st_size} bytes\n")

d = docx.Document(str(out))
all_text = " ".join(p.text for p in d.paragraphs)
all_text += " " + " ".join(c.text for t in d.tables for r in t.rows for c in r.cells)
jinja_leak = "{%" in all_text or "{{" in all_text
sys.stderr.write(f"jinja leak in output: {jinja_leak}\n")
sys.stderr.write(f"tables in output: {len(d.tables)}\n")
for i, t in enumerate(d.tables):
    sys.stderr.write(f"  table {i}: {len(t.rows)} rows, {len(t.rows[0].cells) if t.rows else 0} cols\n")

sys.stderr.write("\n=== Substitution checks ===\n")
checks = {
    "discipline": "Тестовая дисциплина",
    "epigraph": "Тест — это путь к истине.",
    "competency code": "ОК 01",
    "learning_outcome code": "З-1",
    "step 1 stage": "Орг",
    "step 2 means": "проектор",
    "literature_main with marker": "— Учебник",
    "pedagogical_technologies": "ИКТ",
}
for label, snippet in checks.items():
    ok = snippet in all_text
    sys.stderr.write(f"  {label}: {'OK' if ok else 'MISSING'} ({snippet!r})\n")
