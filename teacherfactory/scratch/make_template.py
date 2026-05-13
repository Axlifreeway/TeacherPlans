import docx
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def replace_text_in_doc(doc, replacements):
    for p in doc.paragraphs:
        for old_txt, new_txt in replacements.items():
            if old_txt in p.text:
                p.text = p.text.replace(old_txt, new_txt)
    
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for old_txt, new_txt in replacements.items():
                    if old_txt in cell.text:
                        cell.text = cell.text.replace(old_txt, new_txt)

doc = docx.Document('docs/16.04.26 Алёшкин АА Технологическая карта SQL.docx')

replacements = {
    "ОПЦ.03 Базы данных": "{{ discipline }}",
    "09.01.03 Оператор информационных систем и ресурсов (далее — 09.01.03), группа предоставлена организаторами — студенты 1 курса на базе среднего общего образования, специальность 09.02.07 Информационные системы и программирование, Колледж цифровых и педагогических технологий": "{{ specialty }}, курс {{ course_number }}, группа {{ group_name }} ({{ students_count }} чел.)",
    "Основы SQL: составление запросов к базе данных": "{{ lesson_topic }}",
    "1 ак. час (45 мин.)": "{{ duration }} мин.",
    "Алёшкин Александр Андреевич": "{{ teacher_name }}",
    "Урок закрепления знаний": "{{ lesson_type }}",
    "Лабораторное занятие с элементами мастер-класса": "{{ lesson_kind }}",
    "Формирование базовых умений составления SQL-запросов к реляционной базе данных: выборка данных (SELECT), фильтрация (WHERE), сортировка (ORDER BY), агрегатные функции (COUNT, AVG) и добавление данных (INSERT INTO).": "{{ goal }}",
    "— Повторить понятийный аппарат: «база данных», «таблица», «запись», «поле», «первичный ключ», «SQL», «SELECT», «INSERT INTO», «WHERE», «агрегатная функция»;\n— Объяснить структуру и синтаксис SQL-запроса: обязательные и необязательные предложения, порядок их следования;\n— Научить составлять SELECT-запросы с условием WHERE, сортировкой ORDER BY, агрегатными функциями COUNT и AVG;\n— Сформировать умение добавлять данные в таблицу через INSERT INTO VALUES с соблюдением типов данных и порядка полей.": "{{ task_educational }}",
    "— Развивать аналитическое мышление: умение формулировать задачу выборки данных                          на естественном языке и переводить её в синтаксис SQL;\n— Развивать умения обучающихся анализировать сущность понятийного аппарата через содержание технической документации (ГОСТ);\n— Развивать умение формулировать и подтверждать примерами свою точку зрения, моделировать ситуации;\n— Развивать результаты учебного занятия: знания, умения, навыки (Таблица 1), общие и профессиональные компетенции, предусмотренные ФГОС СПО (Таблица 2).": "{{ task_developmental }}",
    "— Воспитывать общекультурные характеристики личности (коммуникативность, тактичность, ответственность, коллективизм и др.);\n— Формировать ответственное отношение к точности синтаксиса и корректности вносимых в базу данных сведений;\n— Воспитывать профессиональную культуру: проверять запрос до выполнения, не полагаться                    на интуицию.": "{{ task_upbringing }}",
    "— Отработать приём контекстного обучения: постановка учебной задачи как реальной производственной потребности (запрос к базе данных поликлиники);\n— Продемонстрировать метод поэтапного наращивания сложности: синтаксис → условие → сортировка → агрегация → вставка.": "{{ task_pedagogical }}",
}

replace_text_in_doc(doc, replacements)

# Fix Lesson Structure table (last table)
# We need to keep row 0, 1 (headers), replace row 2 with Jinja tags, and delete row 3..8, keeping row 9 (Итого).
# python-docx doesn't have an easy delete row method, we can remove the `<tr>` element.
table = doc.tables[-1]

# Row 2 (index 2) will be our template row
row_tmpl = table.rows[2]
row_tmpl.cells[0].text = "{% tr for step in lesson_structure %}{{ step.number }}"
row_tmpl.cells[1].text = "{{ step.stage }}"
row_tmpl.cells[2].text = "" # Подэтап
row_tmpl.cells[3].text = "{{ step.time }}"
row_tmpl.cells[4].text = "" # Задачи
row_tmpl.cells[5].text = "{{ step.result }}" # Результаты
row_tmpl.cells[6].text = "{{ step.methods }}" # Приёмы
row_tmpl.cells[7].text = "" # Средства
row_tmpl.cells[8].text = "{{ step.teacher }}"
row_tmpl.cells[9].text = "{{ step.student }}"
row_tmpl.cells[10].text = "{% tr endfor %}"

# Delete rows 3, 4, 5, 6, 7, 8
for _ in range(6):
    tr = table.rows[3]._tr
    table._tbl.remove(tr)

doc.save('docs/template_aleshkin.docx')
print("Template saved to docs/template_aleshkin.docx")
