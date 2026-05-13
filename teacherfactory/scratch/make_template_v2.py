import docx
import shutil

# Make a fresh copy to work with
shutil.copy('docs/16.04.26 Алёшкин АА Технологическая карта SQL.docx', 'docs/template_aleshkin.docx')
doc = docx.Document('docs/template_aleshkin.docx')

def replace_flat(replacements):
    for p in doc.paragraphs:
        for old_txt, new_txt in replacements.items():
            if old_txt in p.text:
                p.text = p.text.replace(old_txt, new_txt)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for old_txt, new_txt in replacements.items():
                    if old_txt in cell.text:
                        cell.text = cell.text.replace(old_txt, new_txt)

# 1. Flat replacements
flat_replacements = {
    "ОПЦ.03 Базы данных": "{{ discipline }}",
    "09.01.03 Оператор информационных систем и ресурсов (далее — 09.01.03), группа предоставлена организаторами — студенты 1 курса на базе среднего общего образования, специальность 09.02.07 Информационные системы и программирование, Колледж цифровых и педагогических технологий": "группа {{ group_name }} ({{ students_count }} чел.), специальность {{ specialty }}, курс {{ course_number }}",
    "Основы SQL: составление запросов к базе данных": "{{ lesson_topic }}",
    "1 ак. час (45 мин.)": "{{ duration }} мин.",
    "Алёшкин Александр Андреевич": "{{ teacher_name }}",
    "Урок закрепления знаний": "{{ lesson_type }}",
    "Лабораторное занятие с элементами мастер-класса": "{{ lesson_kind }}",
    "Формирование базовых умений составления SQL-запросов к реляционной базе данных: выборка данных (SELECT), фильтрация (WHERE), сортировка (ORDER BY), агрегатные функции (COUNT, AVG) и добавление данных (INSERT INTO).": "{{ goal }}",
    "Искусство программирования заключается в том, чтобы организовать сложность": "{{ epigraph }}",
    "Дональд Эрвин Кнут": "{{ epigraph_author }}",
    "Компьютер с лицензионным программным обеспечением, браузер (Яндекс.Браузер), электронная тетрадь.": "{{ teaching_means }}",
    "метод контекстного обучения, информационно-коммуникационные технологии (далее — ИКТ), электронная тетрадь.": "{{ pedagogical_technologies }}"
}
replace_flat(flat_replacements)

# 2. Paragraph Lists Replacement (Tasks, Forms, Techniques, Resources)
def replace_para_list(start_marker, end_marker, loop_var):
    start_idx = -1
    end_idx = -1
    for i, p in enumerate(doc.paragraphs):
        if start_marker in p.text:
            start_idx = i
        elif end_marker and end_marker in p.text and start_idx != -1 and end_idx == -1:
            end_idx = i
            break
    
    if start_idx != -1:
        if end_idx == -1:
            end_idx = start_idx + 2 # Default delete 1 paragraph
        
        # We replace the first item with the loop, and delete the rest
        doc.paragraphs[start_idx+1].text = "{% p for item in " + loop_var + " %}— {{ item.text }}{% p endfor %}"
        # Clear text from remaining items to essentially "delete" them (python-docx lacks paragraph deletion)
        for i in range(start_idx+2, end_idx):
            doc.paragraphs[i].text = ""

replace_para_list("Образовательные (дидактические):", "Развивающие:", "task_educational")
replace_para_list("Развивающие:", "Воспитательные:", "task_developmental")
replace_para_list("Воспитательные:", "Перспективные:", "task_upbringing")
replace_para_list("Перспективные:", "Развитие персональных навыков", "task_perspective")
replace_para_list("Развитие персональных навыков", "Формы организации", "task_personal")
replace_para_list("Формы организации учебной деятельности:", "Результаты обучения", "organization_forms")
replace_para_list("Приёмы обучения:", "Методическое обеспечение:", "teaching_techniques")
replace_para_list("Методическое обеспечение:", "Средства обучения:", "methodological_support")

replace_para_list("Используемая основная литература:", "Дополнительная литература:", "lit_main")
replace_para_list("Дополнительная литература:", "Современные профессиональные базы", "lit_add")
replace_para_list("Современные профессиональные базы данных", "Нормативно-правовая документация:", "databases")
replace_para_list("Нормативно-правовая документация:", "Лицензионное и свободно", "normative_docs")
replace_para_list("Лицензионное и свободно распространяемое", "Интернет-ресурсы:", "software")
replace_para_list("Интернет-ресурсы:", "Эпиграф к уроку:", "internet_resources")


# 3. Table 2: Outcomes (Знания, Умения, Навыки)
# The table has many rows. We'll delete all except one data row and put loop tags.
# Actually, it's easier to just take row 2 as the loop row and delete rows 3..9
t_outcomes = doc.tables[2]
row = t_outcomes.rows[2]
row.cells[0].text = "{% tr for out in learning_outcomes %}{{ out.type }}\n{{ out.code }}"
row.cells[1].text = "{{ out.name }}"
row.cells[2].text = "{{ out.indicator }} {% tr endfor %}"
for _ in range(7):
    t_outcomes._tbl.remove(t_outcomes.rows[3]._tr)
# Remove the empty headers "Умения", "Навыки" which were on rows 4 and 7 originally (now moved up or deleted)
# Wait, let's just delete rows 1 as well to simplify
t_outcomes._tbl.remove(t_outcomes.rows[1]._tr)

# 4. Table 3: Competencies
t_comp = doc.tables[3]
row = t_comp.rows[1]
row.cells[0].text = "{% tr for comp in competencies %}{{ comp.code }}"
row.cells[1].text = "{{ comp.name }}"
row.cells[2].text = "{{ comp.indicator }} {% tr endfor %}"
for _ in range(6):
    t_comp._tbl.remove(t_comp.rows[2]._tr)

# 5. Table 4: Interdisciplinary
t_inter = doc.tables[4]
row = t_inter.rows[1]
row.cells[0].text = "{% tr for conn in interdisciplinary_connections %}{{ conn.outcome_source }}"
row.cells[1].text = "{{ conn.discipline }}"
row.cells[2].text = "{{ conn.topic }}"
row.cells[3].text = "{{ conn.outcome_target }}"
row.cells[4].text = "{{ conn.indicator }} {% tr endfor %}"
for _ in range(1):
    t_inter._tbl.remove(t_inter.rows[2]._tr)

# 6. Table 5: Methods
t_methods = doc.tables[5]
t_methods.rows[0].cells[1].text = "{{ method_source }}"
t_methods.rows[1].cells[1].text = "{{ method_character }}"
t_methods.rows[2].cells[1].text = "{{ method_independence }}"
t_methods.rows[3].cells[1].text = "{{ method_health }}"

# 7. Table 6: Lesson Structure
t_struct = doc.tables[-1]
row = t_struct.rows[2]
row.cells[0].text = "{% tr for step in lesson_structure %}{{ step.number }}"
row.cells[1].text = "{{ step.stage }}"
row.cells[2].text = "{{ step.substage }}"
row.cells[3].text = "{{ step.time }}"
row.cells[4].text = "{{ step.tasks }}"
row.cells[5].text = "{{ step.result }}"
row.cells[6].text = "{{ step.methods }}"
row.cells[7].text = "{{ step.means }}"
row.cells[8].text = "{{ step.teacher }}"
row.cells[9].text = "{{ step.student }}"
row.cells[10].text = "{{ step.control }} {% tr endfor %}"
for _ in range(6):
    t_struct._tbl.remove(t_struct.rows[3]._tr)

doc.save('docs/template_aleshkin_v2.docx')
print("V2 Template saved to docs/template_aleshkin_v2.docx")
