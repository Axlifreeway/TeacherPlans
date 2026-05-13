"""
Сборка рабочего шаблона технологической карты урока из docx-эталона Алёшкина
с СОХРАНЕНИЕМ форматирования (шрифты, жирный/курсив, границы таблиц, выравнивание).

Ключевое отличие от v1/v2 — не заменяем `paragraph.text` / `cell.text` целиком
(это сносит все runs и оставляет только дефолтный стиль), а правим текст
ВНУТРИ runs, сохраняя их стилевые атрибуты.

Запуск:
    poetry run python scratch/build_template_v3.py

Выход: docs/template_aleshkin_v2.docx (этот файл прописан в paths.py
       как LESSON_CARD_TEMPLATE).
"""

import copy
import shutil
from pathlib import Path

import docx
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

SRC = Path("docs/16.04.26 Алёшкин АА Технологическая карта SQL.docx")
DST = Path("docs/template_aleshkin_v2.docx")


# ─── Сохраняющие форматирование примитивы ─────────────────────────────────────


def _paragraph_runs_text(p: Paragraph) -> str:
    return "".join(r.text for r in p.runs)


def replace_in_paragraph(p: Paragraph, old: str, new: str) -> bool:
    """
    Заменить подстроку `old` на `new` в параграфе, сохраняя стиль runs.

    Если подстрока укладывается в один run — правим этот run.
    Если она «расползлась» по нескольким — склеиваем все runs в первый
    (теряя промежуточные стили) и пишем `new` туда. Это компромисс: для
    наших целей замена статичных значений и так пишет одинаковым стилем.
    """
    full = _paragraph_runs_text(p)
    if old not in full:
        return False

    # Быстрый путь: подстрока целиком в одном run
    for run in p.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            return True

    # Медленный путь: подстрока на стыке run'ов — собираем в первый run
    if not p.runs:
        # У параграфа нет runs — добавим один
        run = p.add_run(full.replace(old, new))
        return True

    first = p.runs[0]
    first.text = full.replace(old, new)
    for r in p.runs[1:]:
        r.text = ""
    return True


def replace_in_doc(doc, replacements: dict[str, str]) -> None:
    """Запустить replace_in_paragraph по всем параграфам и ячейкам таблиц."""
    for p in doc.paragraphs:
        for old, new in replacements.items():
            if replace_in_paragraph(p, old, new):
                break  # одна замена на параграф — переходим к следующему
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for old, new in replacements.items():
                        if replace_in_paragraph(p, old, new):
                            break


def replace_paragraph_text(p: Paragraph, new: str) -> None:
    """
    Заменить полностью текст параграфа на `new`, сохраняя стиль первого run.

    Используется для строк-маркеров списков (где нам нужна jinja-обёртка
    `{% p for ... %}...{% p endfor %}`).
    """
    if not p.runs:
        p.add_run(new)
        return
    first = p.runs[0]
    first.text = new
    for r in p.runs[1:]:
        r.text = ""


def delete_paragraph(p: Paragraph) -> None:
    """Удалить параграф полностью (python-docx нет API → правим XML)."""
    el = p._element
    el.getparent().remove(el)


# ─── Преобразование списков под jinja-цикл (`{% p for ... %}`) ─────────────────


def replace_para_list_block(
    doc, start_marker: str, end_marker: str | None, loop_var: str
) -> None:
    """
    Между параграфом с `start_marker` и параграфом с `end_marker` лежат
    пункты списка. Сворачиваем их в ОДИН параграф со склейкой через `\\n`:

        [start_marker]                 ← заголовок, не трогаем
        {{ loop_var|join('\\n') }}     ← один параграф со всем списком
        [end_marker]                   ← следующий заголовок

    docxtpl-цикл `{%p for%}` пробовали — пробуксовывает на разных длинах
    списков, поэтому используем простой join. Маркер «— » подставляется
    в to_template_context (см. LessonCard.to_template_context).
    """
    paras = doc.paragraphs
    start = next((i for i, p in enumerate(paras) if start_marker in p.text), -1)
    if start == -1:
        return

    if end_marker is None:
        end = len(paras)
    else:
        end = next(
            (i for i, p in enumerate(paras) if end_marker in p.text and i > start),
            len(paras),
        )

    inner = [p for p in paras[start + 1 : end] if p.text.strip()]
    if not inner:
        return

    # Первый значимый — становится контейнером для всего списка/строки.
    replace_paragraph_text(inner[0], "{{ " + loop_var + " }}")
    # Остальные — удаляем.
    for p in inner[1:]:
        delete_paragraph(p)


# ─── Преобразование таблиц под `{% tr for ... %}` ─────────────────────────────


def set_cell_text_preserve(cell, text: str) -> None:
    """
    Записать `text` в первую paragraph cell, сохраняя стиль первого run
    (а runs других параграфов очистить, лишние параграфы удалить).
    """
    if not cell.paragraphs:
        cell.add_paragraph(text)
        return
    first_p = cell.paragraphs[0]
    replace_paragraph_text(first_p, text)
    for p in cell.paragraphs[1:]:
        delete_paragraph(p)


def remove_row(table, row_index: int) -> None:
    tr = table.rows[row_index]._tr
    table._tbl.remove(tr)


def setup_row_loop(
    table, body_row_index: int, loop_var: str, item_var: str, body_cells: list[str]
) -> None:
    """
    Превратить ряд `body_row_index` в тело docxtpl-цикла `{%tr for ... %}`.

    docxtpl схлопывает <w:tr>, в котором стоит {%tr ...%}, в строку без
    окружающего ряда — поэтому маркеры обязаны быть в *отдельных* рядах:

        [body_row_index]      ← клонируем ↓
        ├── for-marker ряд    ← {%tr for item in xs %} (получится {% for %})
        ├── body ряд          ← реальные {{ item.field }} в каждой ячейке
        └── endfor-marker ряд ← {%tr endfor %} (получится {% endfor %})

    Стиль body-ряда сохраняется, потому что мы клонируем именно его XML.
    """
    body_tr = table.rows[body_row_index]._tr

    # Клонируем body row дважды (для for и endfor маркеров)
    for_tr = copy.deepcopy(body_tr)
    endfor_tr = copy.deepcopy(body_tr)
    body_tr.addprevious(for_tr)
    body_tr.addnext(endfor_tr)

    # Теперь у нас 3 ряда подряд (for, body, endfor) на месте старого body.
    # Перечитаем после вставки.
    for_row = table.rows[body_row_index]
    body_row = table.rows[body_row_index + 1]
    endfor_row = table.rows[body_row_index + 2]

    # for-маркер: только в первой ячейке, остальные — пусто
    set_cell_text_preserve(for_row.cells[0], "{%tr for " + item_var + " in " + loop_var + " %}")
    for c in for_row.cells[1:]:
        set_cell_text_preserve(c, "")

    # body: разнести шаблоны по ячейкам
    for cell, content in zip(body_row.cells, body_cells, strict=False):
        set_cell_text_preserve(cell, content)

    # endfor-маркер: только в первой ячейке
    set_cell_text_preserve(endfor_row.cells[0], "{%tr endfor %}")
    for c in endfor_row.cells[1:]:
        set_cell_text_preserve(c, "")


# ─── Сборка шаблона ───────────────────────────────────────────────────────────


def build() -> None:
    if not SRC.exists():
        raise FileNotFoundError(f"Не найден исходник Эталона: {SRC}")
    shutil.copy(SRC, DST)
    doc = docx.Document(str(DST))

    # 1. Плоские (одна ячейка/параграф = одно значение) замены.
    flat = {
        "ОПЦ.03 Базы данных": "{{ discipline }}",
        "09.01.03 Оператор информационных систем и ресурсов (далее — 09.01.03), группа предоставлена организаторами — студенты 1 курса на базе среднего общего образования, специальность 09.02.07 Информационные системы и программирование, Колледж цифровых и педагогических технологий": "группа {{ group_name }} ({{ students_count }} чел.), специальность {{ specialty }}, курс {{ course_number }}",
        "Основы SQL: составление запросов к базе данных": "{{ lesson_topic }}",
        "1 ак. час (45 мин.)": "{{ duration }} мин.",
        "Алёшкин Александр Андреевич": "{{ teacher_name }}",
        "Урок закрепления знаний": "{{ lesson_type }}",
        "Лабораторное занятие с элементами мастер-класса": "{{ lesson_kind }}",
        "Формирование базовых умений составления SQL-запросов к реляционной базе данных: выборка данных (SELECT), фильтрация (WHERE), сортировка (ORDER BY), агрегатные функции (COUNT, AVG) и добавление данных (INSERT INTO).": "{{ goal }}",
        "Искусство программирования заключается в том, чтобы организовать сложность": "{{ epigraph }}",
        "Дональд Эрвин Кнут": "{{ epigraph_author }}",
        "Компьютер с лицензионным программным обеспечением, браузер (Яндекс.Браузер), электронная тетрадь.": "{{ teaching_means }}",
        "метод контекстного обучения, информационно-коммуникационные технологии (далее — ИКТ), электронная тетрадь.": "{{ pedagogical_technologies }}",
    }
    replace_in_doc(doc, flat)

    # 2. Параграф-списки (4 блока задач + 3 блока методов + 6 блоков ресурсов).
    replace_para_list_block(doc, "Образовательные (дидактические):", "Развивающие:", "task_educational")
    replace_para_list_block(doc, "Развивающие:", "Воспитательные:", "task_developmental")
    replace_para_list_block(doc, "Воспитательные:", "Перспективные:", "task_upbringing")
    replace_para_list_block(doc, "Перспективные:", "Развитие персональных навыков", "task_perspective")
    replace_para_list_block(doc, "Развитие персональных навыков", "Формы организации", "task_personal")

    replace_para_list_block(doc, "Формы организации учебной деятельности:", "Результаты обучения", "organization_forms")
    replace_para_list_block(doc, "Приёмы обучения:", "Методическое обеспечение:", "teaching_techniques")
    replace_para_list_block(doc, "Методическое обеспечение:", "Средства обучения:", "methodological_support")

    replace_para_list_block(doc, "Используемая основная литература:", "Дополнительная литература:", "lit_main")
    replace_para_list_block(doc, "Дополнительная литература:", "Современные профессиональные базы", "lit_add")
    replace_para_list_block(doc, "Современные профессиональные базы данных", "Нормативно-правовая документация:", "databases")
    replace_para_list_block(doc, "Нормативно-правовая документация:", "Лицензионное и свободно", "normative_docs")
    replace_para_list_block(doc, "Лицензионное и свободно распространяемое", "Интернет-ресурсы:", "software")
    replace_para_list_block(doc, "Интернет-ресурсы:", "Эпиграф к уроку:", "internet_resources")

    # 3. Таблица «Результаты обучения» (Знания / Умения / Навыки).
    #    Структура: row 0 — заголовок, row 1 — «Знания» (метка-разделитель),
    #    rows 2.. — примеры. Удаляем все, кроме row 2 (станет body цикла),
    #    и зачищаем row 1 (модель указывает тип в поле out.type).
    t_out = doc.tables[2]
    while len(t_out.rows) > 3:
        remove_row(t_out, 3)
    remove_row(t_out, 1)  # «Знания»
    setup_row_loop(
        t_out,
        body_row_index=1,
        loop_var="learning_outcomes",
        item_var="out",
        body_cells=["{{ out.type }} {{ out.code }}", "{{ out.name }}", "{{ out.indicator }}"],
    )

    # 4. Таблица «Компетенции».
    t_comp = doc.tables[3]
    while len(t_comp.rows) > 2:
        remove_row(t_comp, 2)
    setup_row_loop(
        t_comp,
        body_row_index=1,
        loop_var="competencies",
        item_var="comp",
        body_cells=["{{ comp.code }}", "{{ comp.name }}", "{{ comp.indicator }}"],
    )

    # 5. Таблица «Межпредметные связи».
    t_inter = doc.tables[4]
    while len(t_inter.rows) > 2:
        remove_row(t_inter, 2)
    setup_row_loop(
        t_inter,
        body_row_index=1,
        loop_var="interdisciplinary_connections",
        item_var="conn",
        body_cells=[
            "{{ conn.outcome_source }}",
            "{{ conn.discipline }}",
            "{{ conn.topic }}",
            "{{ conn.outcome_target }}",
            "{{ conn.indicator }}",
        ],
    )

    # 6. Таблица методов обучения (4 строки, фиксированная структура — без цикла).
    t_methods = doc.tables[5]
    set_cell_text_preserve(t_methods.rows[0].cells[1], "{{ method_source }}")
    set_cell_text_preserve(t_methods.rows[1].cells[1], "{{ method_character }}")
    set_cell_text_preserve(t_methods.rows[2].cells[1], "{{ method_independence }}")
    set_cell_text_preserve(t_methods.rows[3].cells[1], "{{ method_health }}")

    # 7. Таблица «Структура учебного занятия» (ход урока).
    #    rows 0-1 — два заголовка (текстовый + нумерация колонок), row 2 — первый этап.
    t_struct = doc.tables[-1]
    while len(t_struct.rows) > 3:
        remove_row(t_struct, 3)
    setup_row_loop(
        t_struct,
        body_row_index=2,
        loop_var="lesson_structure",
        item_var="step",
        body_cells=[
            "{{ step.number }}",
            "{{ step.stage }}",
            "{{ step.substage }}",
            "{{ step.time }}",
            "{{ step.tasks }}",
            "{{ step.result }}",
            "{{ step.methods }}",
            "{{ step.means }}",
            "{{ step.teacher }}",
            "{{ step.student }}",
            "{{ step.control }}",
        ],
    )

    doc.save(str(DST))
    print(f"Шаблон собран: {DST}")


if __name__ == "__main__":
    build()
