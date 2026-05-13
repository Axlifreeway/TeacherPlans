import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from teacherfactory.documents.assessment_materials import ASSESSMENT_MATERIALS
from teacherfactory.model import (
    AssessmentCompetencySummary,
    AssessmentMaterials,
    AssessmentPassportRow,
    AssessmentTestItem,
)
from teacherfactory.render import render_document


def stub() -> AssessmentMaterials:
    return AssessmentMaterials(
        discipline="Математический анализ",
        specialty="09.01.03 Оператор ИС",
        semester=2,
        final_form="дифференцированный зачёт",
        grading_system="Применяется 4-балльная шкала: отлично/хорошо/удовл/неудовл.",
        passport_rows=[
            AssessmentPassportRow(
                topic="Тема 1.1. Действительные числа",
                tool="Тест в количестве 15 заданий",
                competencies="ОК 01\nОК 02",
                criteria="Отлично — 90-100%; хорошо — 71-89%; удовл — 70%; неудовл — <70%.",
            ),
            AssessmentPassportRow(
                topic="Тема 1.2. Предел функции",
                tool="Тест в количестве 15 заданий",
                competencies="ОК 01",
                criteria="Аналогично",
            ),
        ],
        competency_summary=[
            AssessmentCompetencySummary(
                code="ОК 01",
                name="Выбирать способы решения задач",
                items_count=10,
            ),
            AssessmentCompetencySummary(
                code="ОК 02",
                name="Использовать средства поиска",
                items_count=5,
            ),
        ],
        test_items=[
            AssessmentTestItem(
                competency_code="ОК 01",
                number=1,
                instruction_and_task="Выберите. К рациональным относятся: а) -15 б) π в) e г) 40.33",
                answer_key="г) 40.33",
            ),
            AssessmentTestItem(
                competency_code="ОК 02",
                number=2,
                instruction_and_task="Производная sin(x) равна... а) cos x б) -cos x",
                answer_key="а) cos x",
            ),
        ],
    )


out = Path("scratch/assessment_test_output.docx")
render_document(ASSESSMENT_MATERIALS, stub(), out)
sys.stderr.write(f"OK: {out} -> {out.stat().st_size} bytes\n")

import docx
d = docx.Document(str(out))
text = " ".join(p.text for p in d.paragraphs)
text += " " + " ".join(c.text for t in d.tables for r in t.rows for c in r.cells)
sys.stderr.write(f"jinja leak: {('{%' in text) or ('{{' in text)}\n")
for kw in ["Математический анализ", "ОК 01", "Тема 1.1", "40.33", "ВСЕГО"]:
    sys.stderr.write(f"  {kw!r}: {'OK' if kw in text else 'MISS'}\n")
