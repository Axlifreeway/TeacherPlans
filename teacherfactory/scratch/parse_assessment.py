"""Дамп структуры docx-файла в JSON для анализа.

Запуск: poetry run python scratch/parse_assessment.py [path-to.docx]
"""

import json
import sys
from pathlib import Path

import docx

DEFAULT = Path("docs/! ИИКолледж Шаблон_Оценочные материалы.docx")
src = Path(sys.argv[1]) if len(sys.argv) > 1 else DEFAULT
doc = docx.Document(str(src))

out = {
    "paragraphs": [p.text for p in doc.paragraphs if p.text.strip()],
    "tables": [
        [[c.text.strip() for c in row.cells] for row in t.rows] for t in doc.tables
    ],
}
out_path = Path("scratch_assessment.json")
out_path.write_text(json.dumps(out, ensure_ascii=False, indent=2), encoding="utf-8")
print(f"Готово: {out_path} ({len(out['paragraphs'])} параграфов, {len(out['tables'])} таблиц)")
